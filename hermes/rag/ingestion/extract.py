"""Text extraction — port of ``ingestion/text-extraction/extract.ts`` (dispatch by MIME).

Text/markdown/HTML/CSV are handled with the standard library (no new
dependencies). PDF/DOCX/XLSX are extracted via optional libraries that are
*lazy-imported* — mirroring Hermes's ``tools/lazy_deps.py`` philosophy of not
pulling heavy parsers into the base install. When such a library is absent the
extractor raises :class:`ExtractionUnavailable` with the pip name, so the caller
can surface a precise, actionable message instead of a vague failure.

Images would use Gemini vision (``analyzeMedia``) for OCR+caption; that path is
left as a hook (``extract_image``) for the multimodal follow-on.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass


class ExtractionUnavailable(RuntimeError):
    """Raised when a format needs an optional parser that isn't installed."""

    def __init__(self, fmt: str, pip_name: str):
        super().__init__(f"{fmt} extraction requires the '{pip_name}' package (pip install {pip_name})")
        self.fmt = fmt
        self.pip_name = pip_name


@dataclass
class ExtractedDocument:
    modality: str  # 'text' | 'image' | 'video'
    text: str


def _decode_text(buffer: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return buffer.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return buffer.decode("utf-8", errors="replace")


def _html_to_text(html: str) -> str:
    # Drop script/style, convert block tags to newlines, strip remaining tags.
    html = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", html)
    html = re.sub(r"(?i)<(br|/p|/div|/li|/tr|/h[1-6])\s*>", "\n", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    text = (
        text.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
    )
    return re.sub(r"[ \t]+", " ", text)


def _csv_to_text(raw: str) -> str:
    out_lines: list[str] = []
    reader = csv.reader(io.StringIO(raw))
    for row in reader:
        out_lines.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(out_lines)


def _pdf_text(buffer: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise ExtractionUnavailable("PDF", "pypdf") from exc
    reader = PdfReader(io.BytesIO(buffer))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(pages).replace("\x00", "")


def _docx_text(buffer: bytes) -> str:
    try:
        import docx  # type: ignore  # python-docx
    except Exception as exc:  # noqa: BLE001
        raise ExtractionUnavailable("DOCX", "python-docx") from exc
    document = docx.Document(io.BytesIO(buffer))
    return "\n\n".join(p.text for p in document.paragraphs if p.text)


def _xlsx_text(buffer: bytes) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise ExtractionUnavailable("XLSX", "openpyxl") from exc
    wb = load_workbook(io.BytesIO(buffer), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [("" if c is None else str(c)) for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
_XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}


def extract_from_buffer(*, buffer: bytes, mime_type: str, file_name: str = "") -> ExtractedDocument:
    """Dispatch extraction by MIME type (with a filename-extension fallback)."""
    mt = (mime_type or "").lower()
    name = (file_name or "").lower()

    if mt.startswith("image/") or mt.startswith("video/"):
        # Multimodal OCR/caption is a separate (Gemini vision) follow-on.
        return ExtractedDocument(modality="image" if mt.startswith("image/") else "video", text="")

    if mt == "application/pdf" or name.endswith(".pdf"):
        return ExtractedDocument("text", _pdf_text(buffer))
    if mt in _DOCX_MIMES or name.endswith(".docx") or name.endswith(".doc"):
        return ExtractedDocument("text", _docx_text(buffer))
    if mt in _XLSX_MIMES or name.endswith(".xlsx") or name.endswith(".xls"):
        return ExtractedDocument("text", _xlsx_text(buffer))

    raw = _decode_text(buffer)
    if mt in ("text/csv", "text/tsv") or name.endswith(".csv") or name.endswith(".tsv"):
        return ExtractedDocument("text", _csv_to_text(raw))
    if mt in ("text/html", "application/xhtml+xml") or name.endswith(".html") or name.endswith(".htm"):
        return ExtractedDocument("text", _html_to_text(raw))
    return ExtractedDocument("text", raw)
